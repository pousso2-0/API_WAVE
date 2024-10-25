from docx import Document

# Creating a new Document
doc = Document()

# Adding Title
doc.add_heading('Je suis un Développeur Full Stack', level=1)


# Adding Table
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'LE SAVOIR'
hdr_cells[1].text = 'LE SAVOIR-FAIRE'
hdr_cells[2].text = 'LE SAVOIR-ÊTRE'

# Data to populate the table
data = [
    ['Front-end : maîtriser HTML5, CSS3, JavaScript, React, Vue.js', 'Intégration : créer des interfaces responsives, optimiser les performances', 'Créativité : proposer des solutions innovantes'],
    ['Back-end : Python, Node.js, Java, architecture microservices', 'Développement : construire des APIs robustes, gérer le cache', 'Précision : attention aux détails dans le code'],
    ['Cloud : comprendre AWS, Azure, Google Cloud', 'Administration : déployer et maintenir des applications cloud', 'Autonomie : capacité à travailler indépendamment'],
    ['Tests : connaître les types de tests (unitaires, E2E)', 'Testing : écrire et maintenir des suites de tests automatisés', 'Rigueur : assurer la qualité du code'],
    ['CI/CD : maîtriser Git, Jenkins, GitHub Actions', 'Pipeline : automatiser les déploiements, gérer les versions', 'Fiabilité : respect des procédures de déploiement'],
    ['Base de données : MongoDB, PostgreSQL, Redis', 'Optimisation : requêtes performantes, indexation efficace', 'Méthodique : organisation structurée des données'],
    ['Containerisation : Docker, Kubernetes', 'Orchestration : gérer des clusters, scalabilité', 'Persévérance : résolution des problèmes complexes'],
    ['Architecture : Design patterns, SOLID, DRY', 'Conception : créer des solutions évolutives et maintenables', 'Vision globale : comprendre l’impact des décisions techniques'],
    ['Gestion de cache : Redis, Memcached, CDN', 'Performance : optimiser les temps de réponse, gérer la mise en cache', 'Proactivité : anticiper les problèmes de performance'],
    ['Sécurité : OWASP, authentification JWT, OAuth', 'Protection : implémenter des mesures de sécurité robustes', 'Vigilance : maintenir une veille sur les failles de sécurité'],
    ['API Gateway : Kong, AWS API Gateway', 'Routage : gérer les routes, load balancing, rate limiting', 'Organisation : structurer les points d’entrée API'],
    ['Message Brokers : RabbitMQ, Apache Kafka', 'Communication : implémenter des systèmes événementiels', 'Anticipation : prévoir les scénarios de défaillance'],
    ['Monitoring : ELK Stack, Prometheus, Grafana', 'Surveillance : mettre en place des tableaux de bord, alertes', 'Responsabilité : assurer la disponibilité des services'],
    ['Documentation : Swagger, OpenAPI, Markdown', 'Rédaction : documenter le code et les APIs', 'Communication : clarté dans les explications techniques'],
    ['Clean Code : principes SOLID, Clean Architecture', 'Refactoring : améliorer la qualité du code existant', 'Excellence : recherche constante de la qualité'],
    ['SEO : métadonnées, performances, accessibilité', 'Optimisation : améliorer le référencement des applications', 'Conscience : impact du code sur l’expérience utilisateur'],
    ['Analytics : Google Analytics, Mixpanel, Amplitude', 'Analyse : implémenter le tracking, analyser les données utilisateurs', 'Curiosité : comprendre les comportements utilisateurs'],
    ['Progressive Web Apps : Service Workers, Web Workers', 'Développement : créer des applications web progressives', 'Innovation : adopter les nouvelles technologies'],
    ['WebSockets : Socket.io, WebSocket API', 'Temps réel : implémenter des communications bidirectionnelles', 'Adaptabilité : gérer les connexions en temps réel'],
    ['Version Control : Git flow, branches stratégies', 'Collaboration : gérer les conflits, revues de code', 'Esprit d’équipe : partager ses connaissances'],
    ['Tests de charge : JMeter, Artillery', 'Performance : tester la scalabilité des applications', 'Prévoyance : anticiper les pics de charge'],
    ['Debug & Profiling : Chrome DevTools, Node profiler', 'Diagnostic : identifier et résoudre les problèmes', 'Patience : persévérer dans la résolution de bugs'],
    ['Accessibilité : WCAG, ARIA', 'Implémentation : rendre les applications accessibles', 'Empathie : considérer tous les utilisateurs'],
    ['Internationalisation : i18n, l10n', 'Localisation : adapter les applications pour différents marchés', 'Ouverture : sensibilité aux différences culturelles'],
]

# Adding data to table
for row_data in data:
    row = table.add_row().cells
    for i, cell in enumerate(row_data):
        row[i].text = cell

# Save the document
doc.save("/home/ehac6/ehac/dev/js/typescript/app-wave/Competences_Developpeur_Full_Stack.docx")
